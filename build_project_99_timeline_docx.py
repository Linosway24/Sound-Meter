from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Project_99_SLM_timeline_summary.docx"


ENTRIES = [
    (
        "October 20, 2025",
        [
            "USAFSAM-099 / Interactive Equipment Module - SLM was approved.",
            "Target end date was added as tentative.",
            "Project scope recorded: interactive Sound Level Meter module covering calibration, SLM readings, and octave-band analysis surveys.",
        ],
    ),
    (
        "October 21, 2025",
        [
            "Equipment request submitted to loan physical devices to SierTek as reference material for the OED training tools.",
            "Request explained that physical devices were needed to build an accurate digital training platform.",
            "Return date needs verification because the subject and body conflict.",
        ],
    ),
    (
        "October 29, 2025",
        [
            "Michael checked whether the equipment request had been approved and whether equipment was available for pickup.",
            "Cooksey replied that he had pinged the Equipment Cell for a status update.",
        ],
    ),
    (
        "October 30, 2025",
        [
            "Equipment became available for pickup.",
            "Demo and pickup were coordinated for 1400 at E308E.",
            "Demo and equipment pickup occurred.",
            "Primary POCs were identified: SSgt Quinn for SLM and SSgt Kimani for Dosimeter.",
            "SSgt Quinn sent SLM instructions as an attachment.",
        ],
    ),
    (
        "December 22, 2025",
        [
            "Michael reported that all device buttons and associated functionality had been mapped.",
            "Michael was working through SSgt Quinn's documentation and asked to demonstrate the current version or upload it to Blackboard.",
        ],
    ),
    (
        "December 29, 2025",
        [
            "SSgt Quinn requested a Teams review for the next day at 1000 and asked that the current version be uploaded to Blackboard.",
            "Michael confirmed the meeting and asked which Blackboard sandbox to use.",
        ],
    ),
    (
        "December 30, 2025",
        [
            "SSgt Quinn directed upload to SND_BEA_ULTRA.",
            "Michael identified an audio issue specific to Blackboard Ultra and provided an external Articulate Review link for testing.",
            "Cooksey confirmed the link worked on NIPR/Microsoft Edge and that audio worked in that version.",
            "Feedback captured: add zoom/crop, fix calibration screen zoom, increase dB scale from 70 to 140, fix graph/timeline scaling, and review high-dB simulation levels.",
        ],
    ),
    (
        "January 6, 2026",
        [
            "Michael reported Blackboard Ultra audio playback was resolved.",
            "Most requested changes were complete; zoom in/out functionality remained in progress.",
        ],
    ),
    (
        "January 9, 2026",
        [
            "Cooksey said the team would test functionality and wanted another check-in before the next phase.",
            "Cooksey requested a more screen-by-screen storyboard for the scenario portion of the interactive.",
        ],
    ),
    (
        "January 12, 2026",
        [
            "Michael suggested an in-person working session to walk through screens and capture photos/video clips.",
            "Cooksey said USAFSAM wanted to provide a drafted storyboard before an in-person walkthrough.",
            "Cooksey asked to temporarily shift focus to the dosimeter interactive and swap equipment for newer dosimeter models.",
        ],
    ),
    (
        "February 20, 2026",
        [
            "Cooksey contacted the 445th about partnering on hazardous-noise scenario capture.",
            "Possible site-visit windows were discussed, with March 4 PM emerging as the preferred option.",
            "Cooksey confirmed March 4 at 1300 at the hangar bay and described the goal: realistic hazardous-noise survey scenarios with noise-producing equipment/processes and spaces.",
            "Draft SLM storyboard was sent, outlining setup/operation demos, virtual surveys, octave-band analysis, and a hints toggle.",
        ],
    ),
    (
        "February 27, 2026",
        [
            "Cooksey sent meeting notes after the planning meeting.",
            "445th contacts were asked to provide the timeline/order for visiting shops and recording processes.",
        ],
    ),
    (
        "March 4, 2026",
        [
            "Planned 445th site visit / hazardous-noise capture at the hangar bay.",
        ],
    ),
    (
        "March 6, 2026",
        [
            "Cooksey reported all needed noise data had been captured in less time than expected.",
            "Potential remaining need identified: return later for 3D photographs of shops.",
        ],
    ),
    (
        "April 2, 2026",
        [
            "Cooksey requested updates on USAFSAM-099 and USAFSAM-112.",
            "For SLM, he asked to see current progress on the digital noise environments/scenario portion.",
            "Michael replied that the SLM scenario portion was in prototyping, additional photo assets might be needed, test renders were underway, and walkthrough audio/instructional text was being finalized.",
        ],
    ),
    (
        "April 3, 2026",
        [
            "Michael sent voice samples for the SLM interactive and future equipment interactives.",
            "Hanson acknowledged receipt and said he would listen.",
        ],
    ),
    (
        "April 8, 2026",
        [
            "Cooksey requested return of checked-out MultiRAE/SLM equipment.",
            "Michael said he could bring them over the next morning at 1000.",
            "Cooksey selected Eryn as the preferred voice for the interactive.",
        ],
    ),
    (
        "April 9, 2026",
        [
            "Michael reported that the SLM and MultiRAE had been returned to the Equipment Cell.",
            "Cooksey acknowledged the return.",
        ],
    ),
    (
        "April 17, 2026",
        [
            "Cooksey followed up again for updates on USAFSAM-099 and USAFSAM-112.",
        ],
    ),
    (
        "April 27, 2026",
        [
            "Cooksey asked whether a working demo could be provided that week for either the SLM or DRI prototype.",
            "Cooksey requested documentation of the SLM scenario pivot to a static-image approach, including whether the reason was hardware, software, design, time, or another issue.",
            "Cooksey also asked for demos or screenshots of the attempted approach before the pivot.",
        ],
    ),
    (
        "May 8, 2026",
        [
            "Cooksey requested OED project updates and proposed a meeting for May 18.",
        ],
    ),
    (
        "May 21, 2026",
        [
            "Cooksey still wanted a meeting on Project 099 / SLM to walk through the design process, what went wrong, what was discovered, and what needed to shift.",
            "He proposed meeting the following Tuesday afternoon.",
        ],
    ),
    (
        "May 26, 2026",
        [
            "OED project update meeting was coordinated for 1230.",
        ],
    ),
    (
        "June 1, 2026",
        [
            "Michael documented agreed update dates after the project meeting.",
            "Project 99 / SLM update was set for July 2, focused on progress toward a low-fidelity SLM interaction and scenario functionality.",
        ],
    ),
    (
        "June 3, 2026",
        [
            "Cooksey sent USAFSAM-99 SLM Interactive Walkthrough Demo Notes.",
            "DoD SAFE notes were dropped off, a design conversation was suggested for the following week, and Michael said he would review the notes and follow up about a meeting.",
        ],
    ),
    (
        "June 5, 2026",
        [
            "Michael proposed bi-weekly meetings for open OE/BEA projects.",
            "Cooksey agreed and proposed a 1500 meeting to review open projects, deep dives, ideas, and roadblocks.",
        ],
    ),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_fixed_table_width(table, widths_inches):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def add_bullet(cell, text):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.2)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Project 99 / USAFSAM-099 SLM Timeline")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    sr = subtitle.add_run("Major dates and topics only. Compiled from Apple Mail messages and search results on June 8, 2026.")
    sr.font.name = "Calibri"
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(85, 85, 85)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    nr = note.add_run(
        "Note: Entries are summarized for planning/reference use. Some late April-June items came from broader Mail search results and should be verified against full message exports if used formally."
    )
    nr.italic = True
    nr.font.name = "Calibri"
    nr.font.size = Pt(9.5)
    nr.font.color.rgb = RGBColor(85, 85, 85)

    for date, topics in ENTRIES:
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_before = Pt(8)
        date_p.paragraph_format.space_after = Pt(2)
        date_p.paragraph_format.keep_with_next = True
        date_run = date_p.add_run(date)
        date_run.bold = True
        date_run.font.name = "Calibri"
        date_run.font.size = Pt(11.2)
        date_run.font.color.rgb = RGBColor(31, 77, 120)

        for topic in topics:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.26)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(topic)
            run.font.name = "Calibri"
            run.font.size = Pt(10.3)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Project 99 SLM Timeline Summary")
    fr.font.name = "Calibri"
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
